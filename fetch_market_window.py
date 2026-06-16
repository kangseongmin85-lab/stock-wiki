#!/usr/bin/env python3
"""
fetch_market_window.py — 장마감~새벽 시황 데이터 수집 워드 리포트

윈도우:
  - 화~금 마감: 전일 15:30 KST → 당일 05:30 KST  (14시간)
  - 월요일 마감: 금요일 15:30 KST → 월요일 05:30 KST  (62시간, 주말 포함)
  - 토/일 마감: 금요일 15:30 KST → 당일 05:30 KST  (디버그용)

데이터:
  1. Notion DB(종목재료정리)에서 윈도우 시간대 created 페이지 전수 조회
  2. 네이버 검색 API: 활성 테마 키워드별 윈도우 시간대 뉴스

출력: "오늘의 시황/시황_YYYY-MM-DD.docx"

사용법:
  python fetch_market_window.py                       # 오늘 05:30 마감 윈도우
  python fetch_market_window.py --date 2026-05-18     # 지정 날짜 05:30 마감
  python fetch_market_window.py --dry-run             # docx 저장 없이 콘솔 통계만
"""

import os, json, re, sys, time, argparse
import urllib.request, urllib.parse
from datetime import datetime, timedelta, timezone
from pathlib import Path
from email.utils import parsedate_to_datetime
from collections import defaultdict, Counter

# Step 2: 종목 분류·매핑·추출
from stock_taxonomy import (
    load_taxonomy, build_stock_name_index, extract_stocks, map_to_industry,
)

# weekly_db: 시황 아카이브 DB 가 주간 로테이션이라 윈도우에 걸린 모든 DB 조회 필요
sys.path.insert(0, str(Path(__file__).parent))
try:
    from weekly_db import get_db_ids_for_window, get_parent_page_id
    _HAS_WEEKLY = True
except Exception as _e:
    print(f"[경고] weekly_db import 실패 — 단일 DB 모드: {_e}")
    _HAS_WEEKLY = False

# Windows 콘솔 UTF-8
if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ── 설정 ─────────────────────────────────────────────────────────────────
BASE_DIR    = Path(__file__).parent
CONFIG_PATH = BASE_DIR / "config.json"
OUTPUT_DIR  = BASE_DIR / "오늘의 시황"


def _load_config():
    if CONFIG_PATH.exists():
        try:
            return json.loads(CONFIG_PATH.read_text(encoding="utf-8"))
        except Exception:
            return {}
    return {}


_CFG = _load_config()
NOTION_TOKEN     = _CFG.get("NOTION_TOKEN", "")        or os.getenv("NOTION_TOKEN", "")
NOTION_DB_ID     = _CFG.get("NOTION_DB_ID", "")        or os.getenv("NOTION_DB_ID", "")
NAVER_CLIENT_ID  = _CFG.get("NAVER_CLIENT_ID", "")     or os.getenv("NAVER_CLIENT_ID", "")
NAVER_CLIENT_SEC = _CFG.get("NAVER_CLIENT_SECRET", "") or os.getenv("NAVER_CLIENT_SECRET", "")

KST = timezone(timedelta(hours=9))

# 활성 테마 — fetch_news.py 와 동기화 (수정 시 함께 변경)
NAVER_SEARCH_QUERIES = [
    "HBM", "AI데이터센터", "방산", "로봇", "원자력", "원전", "이차전지",
    "자율주행", "휴머노이드", "조선업",
    # 시황·매크로
    "코스피", "코스닥", "환율", "나스닥", "S&P500",
]


# ── 윈도우 계산 ───────────────────────────────────────────────────────────
def compute_window(target_datetime_str=None):
    """
    end_kst: 실행 시각 (= 사용자가 더블클릭한 순간).
             target_datetime_str 지정 시 그 시각으로 고정 (테스트용).
    start_kst: end 보다 과거인 "가장 최근 평일 15:30".
               주말은 건너뜀. 토/일/월 아침 실행 시 자동으로 금요일 15:30 까지 거슬러올라감.

    예시:
      화요일 09:00 실행 → 월 15:30 ~ 화 09:00  (17.5시간)
      월요일 09:00 실행 → 금 15:30 ~ 월 09:00  (65.5시간)
      토요일 12:00 실행 → 금 15:30 ~ 토 12:00  (20.5시간)
      화요일 14:00 실행 → 월 15:30 ~ 화 14:00  (오늘 15:30 아직 안 옴)
      화요일 16:00 실행 → 화 15:30 ~ 화 16:00  (30분)
    """
    if target_datetime_str:
        # 'YYYY-MM-DD' 또는 'YYYY-MM-DD HH:MM' 모두 허용
        try:
            end_kst = datetime.strptime(target_datetime_str, "%Y-%m-%d %H:%M").replace(tzinfo=KST)
        except ValueError:
            end_kst = datetime.strptime(target_datetime_str, "%Y-%m-%d").replace(tzinfo=KST)
    else:
        end_kst = datetime.now(KST)

    # 직전 영업일 15:30 까지 거슬러올라감
    candidate = end_kst.replace(hour=15, minute=30, second=0, microsecond=0)
    # 미래거나(오늘 15:30 아직 안 옴), 주말(토=5, 일=6)이면 하루씩 뒤로
    while candidate >= end_kst or candidate.weekday() >= 5:
        candidate -= timedelta(days=1)
    start_kst = candidate
    return start_kst, end_kst


# ── Notion 조회 ───────────────────────────────────────────────────────────
def _query_single_db(db_id, start_iso, end_iso):
    """단일 Notion DB 윈도우 페이지 전수 조회."""
    results, cursor = [], None
    for _ in range(50):
        body = {
            "page_size": 100,
            "filter": {"and": [
                {"timestamp": "created_time", "created_time": {"on_or_after": start_iso}},
                {"timestamp": "created_time", "created_time": {"before":      end_iso}},
            ]},
            "sorts": [{"timestamp": "created_time", "direction": "ascending"}],
        }
        if cursor:
            body["start_cursor"] = cursor
        req = urllib.request.Request(
            f"https://api.notion.com/v1/databases/{db_id}/query",
            data=json.dumps(body).encode("utf-8"),
            headers={
                "Authorization":  f"Bearer {NOTION_TOKEN}",
                "Notion-Version": "2022-06-28",
                "Content-Type":   "application/json",
            },
            method="POST",
        )
        try:
            with urllib.request.urlopen(req, timeout=30) as r:
                data = json.loads(r.read())
        except Exception as e:
            print(f"  [Notion 조회 오류 db={db_id[:8]}] {e}")
            break
        results.extend(data.get("results", []))
        if not data.get("has_more"):
            break
        cursor = data.get("next_cursor")
    return results


def fetch_notion_articles(start_kst, end_kst):
    """시황 아카이브 DB(주간 로테이션) 에서 윈도우 기사 전수 조회.
    윈도우가 두 주차에 걸치면 양쪽 모두 조회.
    """
    if not (NOTION_TOKEN and NOTION_DB_ID):
        print("  [경고] NOTION_TOKEN/NOTION_DB_ID 없음 — Notion 조회 스킵")
        return []

    start_iso = start_kst.astimezone(timezone.utc).strftime("%Y-%m-%dT%H:%M:%S.000Z")
    end_iso   = end_kst.astimezone(timezone.utc).strftime("%Y-%m-%dT%H:%M:%S.000Z")

    # weekly_db 로 윈도우에 걸친 모든 주차 DB ID 조회
    db_ids = []
    if _HAS_WEEKLY:
        try:
            parent = get_parent_page_id(NOTION_TOKEN, NOTION_DB_ID)
            if parent:
                db_ids = get_db_ids_for_window(NOTION_TOKEN, parent, start_kst, end_kst) or []
        except Exception as e:
            print(f"  [weekly_db 조회 실패 — 단일 DB fallback] {e}")

    if not db_ids:
        db_ids = [NOTION_DB_ID]

    print(f"  대상 DB: {len(db_ids)}개  ({', '.join(d[:8]+'..' for d in db_ids)})")

    all_results = []
    for db_id in db_ids:
        n_before = len(all_results)
        all_results.extend(_query_single_db(db_id, start_iso, end_iso))
        print(f"    db={db_id[:8]}.. → {len(all_results) - n_before}건")
    return all_results


def _prop_text(prop):
    if not prop:
        return ""
    t = prop.get("type", "")
    if t == "title":
        return "".join(p.get("plain_text", "") for p in prop.get("title", []))
    if t == "rich_text":
        return "".join(p.get("plain_text", "") for p in prop.get("rich_text", []))
    if t == "select":
        s = prop.get("select")
        return s.get("name", "") if s else ""
    if t == "multi_select":
        return ", ".join(s.get("name", "") for s in prop.get("multi_select", []))
    if t == "date":
        d = prop.get("date")
        return d.get("start", "") if d else ""
    if t == "url":
        return prop.get("url", "") or ""
    return ""


def parse_notion_page(page):
    """시황 아카이브 DB 와 종목재료정리 DB 컬럼명 모두 호환."""
    props = page.get("properties", {})
    return {
        "title":    _prop_text(props.get("제목")) or _prop_text(props.get("종목명")) or "(제목 없음)",
        "category": _prop_text(props.get("카테고리")) or "기타",
        "channel":  _prop_text(props.get("채널")) or "",
        # 시황 DB: "테마태그" (multi_select) / 종목재료 DB: "관련테마"
        "themes":   _prop_text(props.get("테마태그")) or _prop_text(props.get("관련테마")),
        # 시황 DB: "원문내용" / 종목재료 DB: "오늘 기사 내용"
        "content":  _prop_text(props.get("원문내용")) or _prop_text(props.get("오늘 기사 내용")) or "",
        # 시황 DB: "원문링크" (url) / 종목재료 DB: "링크" (rich_text)
        "link":     _prop_text(props.get("원문링크")) or _prop_text(props.get("링크")) or "",
        "created":  page.get("created_time", ""),
    }


# ── 네이버 검색 (윈도우 정확 필터) ─────────────────────────────────────────
def fetch_naver_window(start_kst, end_kst, queries=None):
    queries = queries or NAVER_SEARCH_QUERIES
    if not (NAVER_CLIENT_ID and NAVER_CLIENT_SEC):
        print("  [경고] 네이버 API 키 없음 — 검색 스킵")
        return []

    start_utc = start_kst.astimezone(timezone.utc)
    end_utc   = end_kst.astimezone(timezone.utc)

    seen_links = set()
    articles = []
    for idx, q in enumerate(queries):
        if idx > 0:
            time.sleep(0.3)  # 네이버 검색 API rate limit (10 req/sec) 회피
        try:
            qenc = urllib.parse.quote(q)
            url  = f"https://openapi.naver.com/v1/search/news.json?query={qenc}&display=50&sort=date"
            req  = urllib.request.Request(url, headers={
                "X-Naver-Client-Id":     NAVER_CLIENT_ID,
                "X-Naver-Client-Secret": NAVER_CLIENT_SEC,
            })
            with urllib.request.urlopen(req, timeout=10) as r:
                data = json.loads(r.read())
        except Exception as e:
            print(f"  [네이버 오류: {q}] {e}")
            continue

        for item in data.get("items", []):
            pub_raw = item.get("pubDate", "").strip()
            if not pub_raw:
                continue
            try:
                pub_dt = parsedate_to_datetime(pub_raw)
            except Exception:
                continue
            if pub_dt.tzinfo is None:
                pub_dt = pub_dt.replace(tzinfo=timezone.utc)
            if not (start_utc <= pub_dt < end_utc):
                continue

            title = re.sub(r"<[^>]+>", "", item.get("title", ""))
            title = (title.replace("&quot;", '"').replace("&amp;", "&")
                          .replace("&lt;", "<").replace("&gt;", ">").replace("&apos;", "'").strip())
            link = item.get("originallink") or item.get("link", "")
            desc = re.sub(r"<[^>]+>", "", item.get("description", "")).strip()[:300]

            if not title or link in seen_links:
                continue
            seen_links.add(link)

            articles.append({
                "query": q,
                "title": title,
                "link":  link,
                "desc":  desc,
                "pub":   pub_dt.astimezone(KST),
            })

    articles.sort(key=lambda a: a["pub"])
    return articles


# ── 워드 리포트 생성 (시그널 리포트 스타일) ────────────────────────────────
def add_hyperlink(paragraph, url, text, color="0066CC", underline=True, bold=False, size=11):
    """python-docx 미지원 — OxmlElement 로 워드 하이퍼링크 직접 추가. 클릭 시 브라우저 오픈."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    if color:
        c = OxmlElement("w:color"); c.set(qn("w:val"), color); rPr.append(c)
    if underline:
        u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    if bold:
        b = OxmlElement("w:b"); rPr.append(b)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(int(size * 2))); rPr.append(sz)
    new_run.append(rPr)

    t = OxmlElement("w:t"); t.text = text; t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def _title_para(doc, prefix: str, title: str, url: str = "", bold: bool = True, size: int = 11):
    """기사 한 줄: prefix(예: '• ') + title 출력. url 있으면 title 을 하이퍼링크화."""
    p = doc.add_paragraph()
    p.add_run(prefix)
    if url:
        add_hyperlink(p, url, title, bold=bold, size=size)
    else:
        r = p.add_run(title)
        r.bold = bold
        r.font.size = Pt(size)
    return p


def _set_meta(run, size=9, color=(0x66, 0x66, 0x66), italic=True):
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(*color)


# 매크로 키워드 (3부에서 별도 표시. 종목 매칭은 일반 기사와 동일하게 처리)
MACRO_QUERIES = {"코스피", "코스닥", "환율", "나스닥", "S&P500"}


def _enrich_articles(notion_pages, naver_articles, taxonomy):
    """각 기사에 매칭된 종목 리스트 추가."""
    names_sorted = build_stock_name_index(taxonomy["stocks"])

    notion_enriched = []
    for page in notion_pages:
        d = parse_notion_page(page)
        text = f"{d['title']} {d['content']}"
        d["stocks"] = sorted(set(extract_stocks(text, names_sorted)))
        d["source"] = "notion"
        notion_enriched.append(d)

    naver_enriched = []
    for art in naver_articles:
        text = f"{art['title']} {art['desc']}"
        art["stocks"] = sorted(set(extract_stocks(text, names_sorted)))
        art["source"] = "naver"
        naver_enriched.append(art)

    return notion_enriched, naver_enriched


def _aggregate(notion_enr, naver_enr, taxonomy):
    """종목별 빈도 + 기사 묶음 + 산업별 그루핑."""
    mentions = Counter()
    articles_by_stock = defaultdict(list)
    for a in notion_enr + naver_enr:
        for s in a["stocks"]:
            mentions[s] += 1
            articles_by_stock[s].append(a)

    industry_to_stocks = defaultdict(list)
    for stock in mentions:
        info = taxonomy["stocks"].get(stock, {})
        ind = map_to_industry(info.get("category", ""), info.get("themes", []))
        industry_to_stocks[ind].append(stock)

    return mentions, articles_by_stock, industry_to_stocks


def build_docx(start_kst, end_kst, notion_pages, naver_articles, output_path, summary_text: str = ""):
    """summary_text 가 있으면 표지 직후에 1페이지 시황 서머리 prepend.
    Opus 가 채팅에서 분석한 결과를 prepend 할 때 사용.
    """
    taxonomy = load_taxonomy()
    notion_enr, naver_enr = _enrich_articles(notion_pages, naver_articles, taxonomy)
    mentions, articles_by_stock, industry_to_stocks = _aggregate(notion_enr, naver_enr, taxonomy)

    doc = Document()

    # ════════ 표지 ════════
    title = doc.add_heading(f"시황 데이터 리포트  {end_kst.strftime('%Y-%m-%d %H:%M')}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    span_h = (end_kst - start_kst).total_seconds() / 3600
    _set_meta(p.add_run(f"윈도우: {start_kst.strftime('%Y-%m-%d %H:%M')} → "
                        f"{end_kst.strftime('%Y-%m-%d %H:%M')} KST ({span_h:.0f}시간)"), size=10)

    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"Notion {len(notion_pages)}건 · 네이버 {len(naver_articles)}건 · {len(mentions)}종목 매칭")
    r2.bold = True

    # 산업 합계 정렬
    ind_summary = sorted(
        industry_to_stocks.items(),
        key=lambda x: sum(mentions[s] for s in x[1]),
        reverse=True,
    )
    doc.add_paragraph()
    p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sm = "산업별 언급: " + "  ·  ".join(
        f"{ind} {sum(mentions[s] for s in stocks)}"
        for ind, stocks in ind_summary[:6]
    )
    _set_meta(p3.add_run(sm), size=10, italic=False)

    doc.add_paragraph("─" * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ════════ Opus 서머리 (있을 때만, 표지 직후) ════════
    if summary_text:
        doc.add_page_break()
        sh = doc.add_heading("🧠 Opus 4.7 시황 서머리 — 검증·분류·추론", level=1)
        sh.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_p = doc.add_paragraph()
        meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_meta(meta_p.add_run("A4 1페이지 분량 · 메모리 행동 룰 적용"), size=9)
        doc.add_paragraph()
        # 서머리 본문 — 줄바꿈을 paragraph 로 분리
        for line in summary_text.split("\n"):
            line = line.rstrip()
            if not line:
                doc.add_paragraph()
                continue
            # 헤더 패턴 (###, ##, **굵게**) 단순 처리
            if line.startswith("### "):
                p = doc.add_paragraph()
                r = p.add_run(line[4:]); r.bold = True; r.font.size = Pt(12)
            elif line.startswith("## "):
                p = doc.add_paragraph()
                r = p.add_run(line[3:]); r.bold = True; r.font.size = Pt(13)
            elif line.startswith("# "):
                p = doc.add_paragraph()
                r = p.add_run(line[2:]); r.bold = True; r.font.size = Pt(14)
            else:
                doc.add_paragraph(line)
        doc.add_page_break()

    # ════════ 1부: 주목 종목 Top 15 ════════
    doc.add_heading("🔥 1부. 주목 종목 Top 15 (언급 빈도순)", level=1)
    doc.add_paragraph("언급 빈도 = 시장 관심도 신호. 동일 종목이 N개 기사에서 언급된 횟수.")

    for rank, (stock, n) in enumerate(mentions.most_common(15), 1):
        info = taxonomy["stocks"].get(stock, {})
        ind = map_to_industry(info.get("category", ""), info.get("themes", []))
        themes = info.get("themes", [])

        p = doc.add_paragraph()
        r = p.add_run(f"{rank}. {stock}"); r.bold = True; r.font.size = Pt(13)
        r2 = p.add_run(f"   {n}건 언급")
        r2.font.size = Pt(11); r2.font.color.rgb = RGBColor(0xCC, 0x00, 0x00); r2.bold = True

        mp = doc.add_paragraph()
        meta = f"   {ind}"
        if info.get("category"): meta += f"  ·  카테고리: {info['category']}"
        if themes: meta += f"  ·  테마: {', '.join(themes[:5])}"
        _set_meta(mp.add_run(meta))

        # 핵심 기사 3건 (제목 하이퍼링크 + 본문 일부)
        for a in articles_by_stock[stock][:3]:
            _title_para(doc, "   • ", a["title"][:110], url=a.get("link", ""), bold=True, size=11)
            src_label = "Notion" if a["source"] == "notion" else "네이버"
            extra = a.get("category", "") if a["source"] == "notion" else a.get("query", "")
            mp = doc.add_paragraph()
            _set_meta(mp.add_run(f"     [{src_label} · {extra}]"))
            body = (a.get("content") or a.get("desc") or "")[:1000]
            if body:
                doc.add_paragraph(f"     {body}")

        if len(articles_by_stock[stock]) > 3:
            ap = doc.add_paragraph()
            _set_meta(ap.add_run(f"   ... 외 {len(articles_by_stock[stock])-3}건 (전체는 2부 참조)"))

    doc.add_page_break()

    # ════════ 2부: 산업별 종목 묶음 ════════
    doc.add_heading("📊 2부. 산업별 종목 묶음", level=1)
    doc.add_paragraph(f"전체 {len(mentions)}종목을 12개 산업으로 분류. 각 산업 안에서 언급 빈도순.")

    for ind, stocks in ind_summary:
        ind_mentions = sum(mentions[s] for s in stocks)
        doc.add_heading(f"< {ind} >   {len(stocks)}종목 · {ind_mentions}건", level=2)

        sorted_stocks = sorted(stocks, key=lambda s: mentions[s], reverse=True)
        for stock in sorted_stocks:
            n = mentions[stock]
            info = taxonomy["stocks"].get(stock, {})
            themes = info.get("themes", [])

            p = doc.add_paragraph()
            r = p.add_run(f"• {stock}"); r.bold = True
            p.add_run(f"  ({n}건)")
            if themes:
                tp = doc.add_paragraph()
                _set_meta(tp.add_run(f"     테마: {', '.join(themes[:4])}"))

            # 기사 제목 묶음 (최대 5건) — 제목 하이퍼링크 + 본문 600자
            for a in articles_by_stock[stock][:5]:
                _title_para(doc, "     - ", a["title"][:120], url=a.get("link", ""), bold=False, size=10)
                body = (a.get("content") or a.get("desc") or "")[:600]
                if body:
                    bp = doc.add_paragraph(f"       {body}")
                    for r in bp.runs:
                        r.font.size = Pt(9)
            if len(articles_by_stock[stock]) > 5:
                ap = doc.add_paragraph()
                _set_meta(ap.add_run(f"     ... 외 {len(articles_by_stock[stock])-5}건"))

    doc.add_page_break()

    # ════════ 3부: 매크로 ════════
    doc.add_heading("🌐 3부. 매크로 (지수·환율·해외)", level=1)
    macro_articles = [a for a in naver_enr if a.get("query") in MACRO_QUERIES]
    if not macro_articles:
        doc.add_paragraph("해당 윈도우에 매크로 키워드 매칭 뉴스 없음.")
    else:
        by_q = defaultdict(list)
        for a in macro_articles:
            by_q[a["query"]].append(a)
        for q in sorted(by_q.keys()):
            arts = by_q[q]
            doc.add_heading(f"[{q}]  {len(arts)}건", level=2)
            for a in arts[:10]:
                _title_para(doc, "• ", a["title"], url=a.get("link", ""), bold=True, size=11)
                mp = doc.add_paragraph()
                _set_meta(mp.add_run(f"   {a['pub'].strftime('%m-%d %H:%M')} KST"))
                if a["desc"]:
                    doc.add_paragraph(f"   {a['desc'][:1000]}")

    # ════════ 4부: 매칭 안 된 기사 (참고용) ════════
    unmatched = [a for a in notion_enr if not a["stocks"]]
    if unmatched:
        doc.add_page_break()
        doc.add_heading(f"📁 4부. 종목 매칭 안 된 Notion 기사 ({len(unmatched)}건)", level=1)
        doc.add_paragraph("참고용. 종목명 추출 실패 또는 시황/매크로/공시 일반 기사.")
        by_cat = defaultdict(list)
        for a in unmatched:
            by_cat[a["category"]].append(a)
        for cat in sorted(by_cat.keys()):
            doc.add_heading(f"[{cat}]  {len(by_cat[cat])}건", level=2)
            for a in by_cat[cat][:5]:
                _title_para(doc, "• ", a["title"][:120], url=a.get("link", ""), bold=True, size=11)
                body = (a.get("content") or "")[:600]
                if body:
                    bp = doc.add_paragraph(f"   {body}")
                    for r in bp.runs:
                        r.font.size = Pt(9)
            if len(by_cat[cat]) > 5:
                p = doc.add_paragraph()
                _set_meta(p.add_run(f"   ... 외 {len(by_cat[cat])-5}건"))

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


# ── 메인 ────────────────────────────────────────────────────────────────────
def main():
    parser = argparse.ArgumentParser(description="장마감~새벽 시황 데이터 수집 워드 리포트")
    parser.add_argument("--datetime", default=None, help="윈도우 끝 시각 (YYYY-MM-DD 또는 \"YYYY-MM-DD HH:MM\"). 미지정시 지금.")
    parser.add_argument("--dry-run", action="store_true", help="docx 저장 없이 콘솔 통계만")
    args = parser.parse_args()

    start_kst, end_kst = compute_window(args.datetime)
    span_h = (end_kst - start_kst).total_seconds() / 3600
    print("=" * 60)
    print(f"  시황 데이터 수집")
    print(f"  윈도우: {start_kst.strftime('%Y-%m-%d %H:%M')} → "
          f"{end_kst.strftime('%Y-%m-%d %H:%M')} KST  ({span_h:.0f}시간)")
    print("=" * 60)

    print(f"\n[1/2] Notion DB 조회...")
    notion_pages = fetch_notion_articles(start_kst, end_kst)
    print(f"  → {len(notion_pages)}건 수집")

    print(f"\n[2/2] 네이버 검색 API ({len(NAVER_SEARCH_QUERIES)}개 키워드)...")
    naver_articles = fetch_naver_window(start_kst, end_kst)
    print(f"  → {len(naver_articles)}건 수집 (윈도우 필터 후)")

    if args.dry_run:
        print(f"\n[dry-run] 저장 생략")
        return

    output = OUTPUT_DIR / f"시황_{end_kst.strftime('%Y-%m-%d_%H%M')}.docx"
    build_docx(start_kst, end_kst, notion_pages, naver_articles, output)
    print(f"\n[OK] 저장 완료: {output}")
    print(f"\n다음 단계: 워드파일을 Claude 대화에 첨부하고 '시황 분석해줘' 라고 요청하세요.")


if __name__ == "__main__":
    main()
